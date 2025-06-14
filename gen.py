"""
This application automates the literature review process by:
1. Using ChatGPT to create search queries based on research questions
2. Suggesting publication corpuses based on keywords
3. Searching and retrieving article metadata and content
4. Facilitating article selection for inclusion in the review
5. Generating a structured literature review using selected methodology
6. Providing an integrated document editor with export capabilities

Built with PySide2 and QML
"""

## Explore a RAG-based approach 

import sys
import os
import json
import re
import requests
import openai
from datetime import datetime
from PyPDF2 import PdfReader
import docx
from docx import Document
from docx.shared import Pt, Inches
import threading
import logging

from PySide2.QtCore import QObject, Signal, Slot, Property, QUrl, Qt, QStringListModel
from PySide2.QtGui import QGuiApplication
from PySide2.QtQml import QQmlApplicationEngine, qmlRegisterType

# Configure logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                    filename='autolitreview.log')
logger = logging.getLogger(__name__)

# Configuration settings
CONFIG_FILE = "config.json"

DEFAULT_CONFIG = {
    "api_key": "",
    "model": "gpt-4o",
    "max_tokens": 4000,
    "temperature": 0.7,
    "publication_databases": [
        {"name": "arXiv", "enabled": True, "url": "https://arxiv.org/search/"},
        {"name": "PubMed", "enabled": True, "url": "https://pubmed.ncbi.nlm.nih.gov/"},
        {"name": "IEEE Xplore", "enabled": True, "url": "https://ieeexplore.ieee.org/search/"},
        {"name": "ACM Digital Library", "enabled": True, "url": "https://dl.acm.org/action/doSearch"},
        {"name": "ScienceDirect", "enabled": True, "url": "https://www.sciencedirect.com/search"}
    ],
    "review_methodologies": [
        "Systematic Review",
        "Meta-analysis",
        "Narrative Review",
        "Scoping Review",
        "Integrative Review"
    ],
    "recent_projects": []
}

class ConfigManager:
    def __init__(self):
        self.config = self.load_config()
    
    def load_config(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r') as f:
                    return json.load(f)
            except Exception as e:
                logger.error(f"Error loading config: {e}")
                return DEFAULT_CONFIG.copy()
        else:
            return DEFAULT_CONFIG.copy()
    
    def save_config(self):
        try:
            with open(CONFIG_FILE, 'w') as f:
                json.dump(self.config, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving config: {e}")
    
    def get_api_key(self):
        return self.config.get("api_key", "")
    
    def set_api_key(self, api_key):
        self.config["api_key"] = api_key
        self.save_config()
    
    def add_recent_project(self, project_path):
        """Add a project to recent projects list"""
        recent = self.config.get("recent_projects", [])
        if project_path in recent:
            recent.remove(project_path)
        recent.insert(0, project_path)
        self.config["recent_projects"] = recent[:10]  # Keep only 10 most recent
        self.save_config()

class Article:
    def __init__(self, title="", authors=None, journal="", year="", doi="", abstract="", 
                 conclusion="", full_text="", url="", source_db=""):
        self.title = title
        self.authors = authors or []
        self.journal = journal
        self.year = year
        self.doi = doi
        self.abstract = abstract
        self.conclusion = conclusion
        self.full_text = full_text
        self.url = url
        self.source_db = source_db
        self.selected = False
        self.notes = ""
        self.local_file_path = ""
    
    def to_dict(self):
        return {
            "title": self.title,
            "authors": self.authors,
            "journal": self.journal,
            "year": self.year,
            "doi": self.doi,
            "abstract": self.abstract,
            "conclusion": self.conclusion,
            "full_text": self.full_text,
            "url": self.url,
            "source_db": self.source_db,
            "selected": self.selected,
            "notes": self.notes,
            "local_file_path": self.local_file_path
        }
    
    @classmethod
    def from_dict(cls, data):
        article = cls()
        for key, value in data.items():
            setattr(article, key, value)
        return article
    
    def extract_conclusion(self):
        """Attempt to extract conclusion from full text if available"""
        if not self.full_text:
            return
        
        # Simple pattern matching for conclusion section
        patterns = [
            r'(?i)conclusion[s]?\s*(.*?)(?:\n\s*[A-Z][a-z]+|\Z)',
            r'(?i)discussion and conclusion[s]?\s*(.*?)(?:\n\s*[A-Z][a-z]+|\Z)',
            r'(?i)summary\s*(.*?)(?:\n\s*[A-Z][a-z]+|\Z)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, self.full_text, re.DOTALL)
            if match:
                self.conclusion = match.group(1).strip()
                return

class Project:
    def __init__(self, name="", path=""):
        self.name = name
        self.path = path
        self.research_questions = []
        self.search_queries = []
        self.selected_databases = []
        self.articles = []
        self.selected_articles = []
        self.review_methodology = ""
        self.review_content = ""
        self.date_created = datetime.now().isoformat()
        self.date_modified = self.date_created
    
    def to_dict(self):
        return {
            "name": self.name,
            "path": self.path,
            "research_questions": self.research_questions,
            "search_queries": self.search_queries,
            "selected_databases": self.selected_databases,
            "articles": [article.to_dict() for article in self.articles],
            "selected_articles": [article.to_dict() for article in self.selected_articles],
            "review_methodology": self.review_methodology,
            "review_content": self.review_content,
            "date_created": self.date_created,
            "date_modified": datetime.now().isoformat()
        }
    
    @classmethod
    def from_dict(cls, data):
        project = cls(name=data.get("name", ""), path=data.get("path", ""))
        project.research_questions = data.get("research_questions", [])
        project.search_queries = data.get("search_queries", [])
        project.selected_databases = data.get("selected_databases", [])
        project.articles = [Article.from_dict(article_data) for article_data in data.get("articles", [])]
        project.selected_articles = [Article.from_dict(article_data) for article_data in data.get("selected_articles", [])]
        project.review_methodology = data.get("review_methodology", "")
        project.review_content = data.get("review_content", "")
        project.date_created = data.get("date_created", datetime.now().isoformat())
        project.date_modified = data.get("date_modified", datetime.now().isoformat())
        return project

    def save(self):
        """Save project to its file path"""
        if not self.path:
            raise ValueError("Project path not set")
        
        self.date_modified = datetime.now().isoformat()
        
        try:
            with open(self.path, 'w') as f:
                json.dump(self.to_dict(), f, indent=2)
        except Exception as e:
            logger.error(f"Error saving project: {e}")
            raise
    
    @classmethod
    def load(cls, path):
        """Load project from file path"""
        try:
            with open(path, 'r') as f:
                data = json.load(f)
            return cls.from_dict(data)
        except Exception as e:
            logger.error(f"Error loading project: {e}")
            raise

class ChatGPTService:
    def __init__(self, config_manager):
        self.config_manager = config_manager
        self.api_key = config_manager.get_api_key()
        self.setup_client()
    
    def setup_client(self):
        if self.api_key:
            openai.api_key = self.api_key
        else:
            logger.warning("OpenAI API key not set")
    
    def update_api_key(self, api_key):
        self.api_key = api_key
        self.config_manager.set_api_key(api_key)
        self.setup_client()
    
    def generate_search_queries(self, research_questions):
        """Generate search queries from research questions using ChatGPT"""
        if not self.api_key:
            raise ValueError("API key not set")
        
        prompt = f"""
        I'm conducting a literature review with the following research questions:
        
        {research_questions}
        
        Please generate 3-5 effective search queries for academic databases that would help me find relevant literature.
        Format each query as a set of terms with appropriate Boolean operators (AND, OR, NOT).
        For each query, provide a brief explanation of what aspect of my research it targets.
        Format your response as a JSON list of objects, each with 'query' and 'explanation' fields.
        """
        
        try:
            response = openai.ChatCompletion.create(
                model=self.config_manager.config.get("model", "gpt-4o"),
                messages=[{"role": "user", "content": prompt}],
                temperature=self.config_manager.config.get("temperature", 0.7),
                max_tokens=self.config_manager.config.get("max_tokens", 4000)
            )
            
            content = response.choices[0].message.content
            # Extract JSON from the response
            json_match = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
            if json_match:
                content = json_match.group(1)
            
            try:
                return json.loads(content)
            except:
                # If parsing fails, try to extract a JSON array
                json_array_match = re.search(r'\[\s*{.*}\s*\]', content, re.DOTALL)
                if json_array_match:
                    return json.loads(json_array_match.group(0))
                else:
                    logger.error(f"Failed to parse ChatGPT response as JSON: {content}")
                    return []
                    
        except Exception as e:
            logger.error(f"Error generating search queries: {e}")
            raise
    
    def suggest_publication_databases(self, research_questions, search_queries):
        """Suggest relevant publication databases based on research questions and search queries"""
        if not self.api_key:
            raise ValueError("API key not set")
        
        all_databases = self.config_manager.config.get("publication_databases", [])
        db_names = [db["name"] for db in all_databases]
        
        prompt = f"""
        I'm conducting a literature review with the following research questions:
        
        {research_questions}
        
        And these search queries:
        
        {search_queries}
        
        From this list of academic databases:
        {", ".join(db_names)}
        
        Please recommend which ones would be most relevant for my research, and explain why.
        Format your response as a JSON list of objects, each with 'database' and 'reason' fields.
        Only include databases from the list provided.
        """
        
        try:
            response = openai.ChatCompletion.create(
                model=self.config_manager.config.get("model", "gpt-4o"),
                messages=[{"role": "user", "content": prompt}],
                temperature=self.config_manager.config.get("temperature", 0.7),
                max_tokens=self.config_manager.config.get("max_tokens", 4000)
            )
            
            content = response.choices[0].message.content
            # Extract JSON from the response
            json_match = re.search(r'```json\s*(.*?)\s*```', content, re.DOTALL)
            if json_match:
                content = json_match.group(1)
            
            try:
                results = json.loads(content)
                # Filter to ensure only valid databases are included
                return [item for item in results if item.get("database") in db_names]
            except:
                logger.error(f"Failed to parse ChatGPT response as JSON: {content}")
                return []
                    
        except Exception as e:
            logger.error(f"Error suggesting databases: {e}")
            raise
    
    def generate_literature_review(self, research_questions, articles, methodology):
        """Generate a literature review based on selected articles and methodology"""
        if not self.api_key:
            raise ValueError("API key not set")
        
        # Prepare article data
        article_data = []
        for article in articles:
            article_info = {
                "title": article.title,
                "authors": article.authors,
                "journal": article.journal,
                "year": article.year,
                "abstract": article.abstract,
                "conclusion": article.conclusion
            }
            article_data.append(article_info)
        
        prompt = f"""
        I'm writing a {methodology} literature review addressing these research questions:
        
        {research_questions}
        
        Here are the selected articles for review:
        
        {json.dumps(article_data, indent=2)}
        
        Please generate a comprehensive literature review following the {methodology} approach.
        
        Include:
        1. Introduction with research questions
        2. Methodology section explaining the review process
        3. Thematic analysis of the literature
        4. Discussion of findings in relation to research questions
        5. Gaps in the literature and suggestions for future research
        6. Conclusion
        7. Suggestions for where figures should be included (marked as [FIGURE: Description of suggested figure])
        
        Format the review with appropriate headings and subheadings.
        """
        
        try:
            response = openai.ChatCompletion.create(
                model=self.config_manager.config.get("model", "gpt-4o"),
                messages=[{"role": "user", "content": prompt}],
                temperature=self.config_manager.config.get("temperature", 0.7),
                max_tokens=self.config_manager.config.get("max_tokens", 8000)
            )
            
            return response.choices[0].message.content
                    
        except Exception as e:
            logger.error(f"Error generating literature review: {e}")
            raise
    
    def expand_review_section(self, review_content, section_title, section_content):
        """Expand a specific section of the literature review"""
        if not self.api_key:
            raise ValueError("API key not set")
        
        prompt = f"""
        I have a literature review with the following section that needs expansion:
        
        # {section_title}
        {section_content}
        
        Please provide a more detailed and comprehensive version of this section, 
        maintaining the same overall structure but adding more depth, analysis, and connections 
        between the ideas presented. Keep the same academic tone and style.
        """
        
        try:
            response = openai.ChatCompletion.create(
                model=self.config_manager.config.get("model", "gpt-4o"),
                messages=[{"role": "user", "content": prompt}],
                temperature=self.config_manager.config.get("temperature", 0.7),
                max_tokens=self.config_manager.config.get("max_tokens", 4000)
            )
            
            return response.choices[0].message.content
                    
        except Exception as e:
            logger.error(f"Error expanding review section: {e}")
            raise

class ArticleRetriever:
    """Handles searching for and retrieving articles from academic databases"""
    
    def __init__(self, config_manager):
        self.config_manager = config_manager
    
    def search_articles(self, query, databases):
        """Search for articles using the given query across specified databases
        
        This is a simplified implementation - in a real application, this would
        connect to database APIs or use web scraping techniques.
        """
        results = []
        
        # Example implementation that would be replaced with actual API calls
        # In a real implementation, you would handle authentication, pagination, rate limits, etc.
        for db in databases:
            try:
                # Simulating API call or web scraping
                logger.info(f"Searching {db} for: {query}")
                
                # In a real implementation, this would make actual requests
                # For example, for arXiv:
                # if db == "arXiv":
                #     results.extend(self._search_arxiv(query))
                
                # For demonstration, create some dummy results
                db_results = self._generate_dummy_results(db, query, 5)
                results.extend(db_results)
                
            except Exception as e:
                logger.error(f"Error searching {db}: {e}")
        
        return results
    
    def _generate_dummy_results(self, database, query, count=5):
        """Generate dummy article results for demonstration purposes"""
        results = []
        keywords = re.findall(r'\b\w+\b', query)
        keywords = [k for k in keywords if k.lower() not in ('and', 'or', 'not')]
        
        for i in range(1, count + 1):
            keyword = keywords[i % len(keywords)] if keywords else "research"
            year = 2020 + (i % 5)
            
            article = Article(
                title=f"Research on {keyword.title()} in Modern Context ({i})",
                authors=[f"Author {i}A", f"Author {i}B"],
                journal=f"Journal of {database} Research",
                year=str(year),
                doi=f"10.1234/jxyz.{year}.{i:04d}",
                abstract=f"This study examines {keyword} in various contexts. "
                         f"We explored how {keyword} affects outcomes and presents "
                         f"a new framework for understanding its implications.",
                url=f"https://example.org/{database.lower()}/{year}/{i:04d}",
                source_db=database
            )
            results.append(article)
        
        return results
    
    def _search_arxiv(self, query):
        """Real implementation for searching arXiv"""
        # This would use arXiv API
        # Example: https://arxiv.org/help/api/user-manual
        pass
    
    def _search_pubmed(self, query):
        """Real implementation for searching PubMed"""
        # This would use NCBI E-utilities
        # Example: https://www.ncbi.nlm.nih.gov/books/NBK25500/
        pass
    
    def retrieve_full_text(self, article):
        """Attempt to retrieve full text for an article
        
        In a real implementation, this would:
        1. Check for open access versions
        2. Use institutional access if available
        3. Handle authentication for subscription services
        4. Parse PDF/HTML to extract text
        """
        logger.info(f"Attempting to retrieve full text for: {article.title}")
        
        # Simulated implementation - would be replaced with real retrieval logic
        try:
            # For demonstration purposes
            article.full_text = (
                f"INTRODUCTION\n\n"
                f"This paper explores {article.title.lower()}. "
                f"Our research was motivated by the need to understand this topic better.\n\n"
                f"METHODOLOGY\n\n"
                f"We conducted a study using qualitative and quantitative methods.\n\n"
                f"RESULTS\n\n"
                f"Our findings indicate significant correlations between variables.\n\n"
                f"DISCUSSION\n\n"
                f"These results suggest important implications for the field.\n\n"
                f"CONCLUSION\n\n"
                f"In conclusion, we have demonstrated that our approach provides "
                f"valuable insights into {article.title.lower()}. Future research "
                f"should focus on expanding these findings."
            )
            
            # Extract conclusion
            article.extract_conclusion()
            
            return True
        except Exception as e:
            logger.error(f"Error retrieving full text: {e}")
            return False
    
    def parse_pdf(self, file_path):
        """Parse a PDF file to extract text and metadata"""
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                
                # Extract text from all pages
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                # Extract metadata
                metadata = reader.metadata
                
                return {
                    "text": text,
                    "metadata": metadata
                }
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return None

class DocumentExporter:
    """Handles exporting the literature review to various document formats"""
    
    def export_docx(self, content, output_path):
        """Export the literature review content to a DOCX file"""
        try:
            doc = Document()
            
            # Set document styles
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # Parse markdown-like content
            lines = content.split('\n')
            in_list = False
            
            for line in lines:
                # Handle headings
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                # Handle lists
                elif line.strip().startswith('- '):
                    if not in_list:
                        in_list = True
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip().startswith('1. ') or line.strip().startswith('2. '):
                    if not in_list:
                        in_list = True
                    # Extract the number and text
                    parts = line.strip().split('. ', 1)
                    if len(parts) > 1:
                        doc.add_paragraph(parts[1], style='List Number')
                # Handle figure placeholders
                elif '[FIGURE:' in line:
                    figure_text = line[line.find('[FIGURE:'):line.find(']')+1]
                    p = doc.add_paragraph()
                    p.add_run(figure_text).italic = True
                    # Add paragraph before and after for spacing
                    doc.add_paragraph()
                # Regular paragraph
                elif line.strip():
                    if in_list:
                        in_list = False
                    doc.add_paragraph(line)
                # Blank line
                else:
                    if not in_list:
                        doc.add_paragraph()
            
            # Save the document
            doc.save(output_path)
            return True
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            return False
    
    def export_text(self, content, output_path):
        """Export the literature review content to a plain text file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            logger.error(f"Error exporting to text: {e}")
            return False
    
    def export_markdown(self, content, output_path):
        """Export the literature review content to a markdown file"""
        try:
            # Content is already in markdown-like format
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            logger.error(f"Error exporting to markdown: {e}")
            return False

# QML Models and Controllers
class ProjectController(QObject):
    projectLoaded = Signal()
    projectSaved = Signal()
    researchQuestionsChanged = Signal()
    searchQueriesChanged = Signal()
    articlesChanged = Signal()
    reviewContentChanged = Signal()
    errorOccurred = Signal(str)
    
    def __init__(self, config_manager, chatgpt_service, article_retriever, doc_exporter):
        super().__init__()
        self.config_manager = config_manager
        self.chatgpt_service = chatgpt_service
        self.article_retriever = article_retriever
        self.doc_exporter = doc_exporter
        self.current_project = Project()
        self._is_loading = False
    
    @Slot(str, str)
    def createNewProject(self, name, path):
        try:
            self.current_project = Project(name=name, path=path)
            self.current_project.save()
            self.config_manager.add_recent_project(path)
            self.projectLoaded.emit()
        except Exception as e:
            logger.error(f"Error creating project: {e}")
            self.errorOccurred.emit(f"Failed to create project: {str(e)}")
    
    @Slot(str)
    def loadProject(self, path):
        try:
            self._is_loading = True
            self.current_project = Project.load(path)
            self.config_manager.add_recent_project(path)
            self._is_loading = False
            self.projectLoaded.emit()
        except Exception as e:
            self._is_loading = False
            logger.error(f"Error loading project: {e}")
            self.errorOccurred.emit(f"Failed to load project: {str(e)}")
    
    @Slot()
    def saveProject(self):
        try:
            self.current_project.save()
            self.projectSaved.emit()
        except Exception as e:
            logger.error(f"Error saving project: {e}")
            self.errorOccurred.emit(f"Failed to save project: {str(e)}")
    
    @Slot(str)
    def addResearchQuestion(self, question):
        self.current_project.research_questions.append(question)
        self.saveProject()
        self.researchQuestionsChanged.emit()
    
    @Slot(int)
    def removeResearchQuestion(self, index):
        if 0 <= index < len(self.current_project.research_questions):
            del self.current_project.research_questions[index]
            self.saveProject()
            self.researchQuestionsChanged.emit()
    
    @Slot(result=str)
    def getResearchQuestionsJson(self):
        return json.dumps(self.current_project.research_questions)
    
    @Slot()
    def generateSearchQueries(self):
        try:
            questions = "\n".join(self.current_project.research_questions)
            queries = self.chatgpt_service.generate_search_queries(questions)
            self.current_project.search_queries = queries
            self.saveProject()
            self.searchQueriesChanged.emit()
        except Exception as e:
            logger.error(f"Error generating search queries: {e}")
            self.errorOccurred.emit(f"Failed to generate search queries: {str(e)}")
    
    @Slot(str)
    def addSearchQuery(self, query_json):
        try:
            query = json.loads(query_json)
            self.current_project.search_queries.append(query)
            self.saveProject()
            self.searchQueriesChanged.emit()
        except Exception as e:
            logger.error(f"Error adding search query: {e}")
            self.errorOccurred.emit(f"Failed to add search query: {str(e)}")
    
    @Slot(int)
    def removeSearchQuery(self, index):
        if 0 <= index < len(self.current_project.search_queries):
            del self.current_project.search_queries[index]
            self.saveProject()
            self.searchQueriesChanged.emit()
    
    @Slot(result=str)
    def getSearchQueriesJson(self):
        return json.dumps(self.current_project.search_queries)
    
    @Slot()
    def suggestDatabases(self):
        try:
            questions = "\n".join(self.current_project.research_questions)
            queries = json.dumps(self.current_project.search_queries)
            suggested_dbs = self.chatgpt_service.suggest_publication_databases(questions, queries)
            
            # Update selected databases
            self.current_project.selected_databases = [db.get("database") for db in suggested_dbs]
            self.saveProject()
            return json.dumps(suggested_dbs)
        except Exception as e:
            logger.error(f"Error suggesting databases: {e}")
            self.errorOccurred.emit(f"Failed to suggest databases: {str(e)}")
            return "[]"
    
    @Slot(str, bool)
    def setDatabaseSelected(self, database, selected):
        if selected and database not in self.current_project.selected_databases:
            self.current_project.selected_databases.append(database)
        elif not selected and database in self.current_project.selected_databases:
            self.current_project.selected_databases.remove(database)
        self.saveProject()
    
    @Slot(result=str)
    def getSelectedDatabasesJson(self):
        return json.dumps(self.current_project.selected_databases)
    
    @Slot(str)
    def searchArticles(self, query_index):
        try:
            index = int(query_index)
            if 0 <= index < len(self.current_project.search_queries):
                query = self.current_project.search_queries[index]["query"]
                results = self.article_retriever.search_articles(query, self.current_project.selected_databases)
                
                # Add new articles to the project
                for article in results:
                    # Check if article already exists (by DOI or title)
                    exists = False
                    for existing in self.current_project.articles:
                        if article.doi and article.doi == existing.doi:
                            exists = True
                            break
                        elif article.title