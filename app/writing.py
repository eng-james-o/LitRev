import docx
from docx import Document
from docx.shared import Pt, Inches
import logging
from .config import logger

class DocumentExporter:
    """Handles exporting the literature review to various document formats"""

    def export_docx(self, content, output_path):
        """Export the literature review content to a DOCX file"""
        try:
            doc = Document()

            # Set document styles
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)

            # Parse markdown-like content
            lines = content.split('\n')
            in_list = False

            for line in lines:
                # Handle headings
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                # Handle lists
                elif line.strip().startswith('- '):
                    if not in_list:
                        in_list = True
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip().startswith('1. ') or line.strip().startswith('2. '):
                    if not in_list:
                        in_list = True
                    # Extract the number and text
                    parts = line.strip().split('. ', 1)
                    if len(parts) > 1:
                        doc.add_paragraph(parts[1], style='List Number')
                # Handle figure placeholders
                elif '[FIGURE:' in line:
                    figure_text = line[line.find('[FIGURE:'):line.find(']')+1]
                    p = doc.add_paragraph()
                    p.add_run(figure_text).italic = True
                    # Add paragraph before and after for spacing
                    doc.add_paragraph()
                # Regular paragraph
                elif line.strip():
                    if in_list:
                        in_list = False
                    doc.add_paragraph(line)
                # Blank line
                else:
                    if not in_list:
                        doc.add_paragraph()

            # Save the document
            doc.save(output_path)
            return True
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            return False

    def export_text(self, content, output_path):
        """Export the literature review content to a plain text file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            logger.error(f"Error exporting to text: {e}")
            return False

    def export_markdown(self, content, output_path):
        """Export the literature review content to a markdown file"""
        try:
            # Content is already in markdown-like format
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            logger.error(f"Error exporting to markdown: {e}")
            return False
